"""
Generator de Rapoarte Word pentru Fisa Pacientului
Foloseste python-docx-template (docxtpl) pentru a genera rapoarte formatate
din datele JSON structurate extrase din transcripții medicale.
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class MedicalReportGenerator:
    """Generator de rapoarte Word pentru fisa pacientului"""

    def __init__(self, template_path: str = None):
        """
        Inițializeaza generatorul de rapoarte

        Args:
            template_path: Calea catre sablonul Word (opțional)
        """
        self.template_path = template_path

    def load_json_data(self, json_path: str) -> Dict[str, Any]:
        """incarca datele din fisierul JSON"""
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def create_simple_report_without_template(self, json_path: str, output_path: str):
        """
        Creeaza un raport Word simplu fara sablon (folosind python-docx)
        Util cand nu ai un sablon predefinit
        """
        # incarca datele
        data = self.load_json_data(json_path)

        # Creeaza un document nou
        doc = Document()

        # Adauga header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "RAPORT MEDICAL - ECOGRAFIE CARDIACa"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Titlu
        title = doc.add_heading('FIsA PACIENTULUI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informații generale
        doc.add_heading('Informații Generale', level=1)
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Light Grid Accent 1'

        # Date raport
        info_table.rows[0].cells[0].text = 'Data raportului:'
        info_table.rows[0].cells[1].text = datetime.now().strftime("%d.%m.%Y %H:%M")
        info_table.rows[1].cells[0].text = 'Tip investigație:'
        info_table.rows[1].cells[1].text = 'Ecografie cardiaca'

        # Masuratori ecografice
        doc.add_heading('Masuratori Ecografice', level=1)

        if data.get('masuratori_ecografice'):
            # Creeaza tabel pentru masuratori
            masuratori_table = doc.add_table(rows=1, cols=3)
            masuratori_table.style = 'Light Grid Accent 1'

            # Header tabel
            hdr_cells = masuratori_table.rows[0].cells
            hdr_cells[0].text = 'Structura Anatomica'
            hdr_cells[1].text = 'Valoare'
            hdr_cells[2].text = 'Unitate'

            # Adauga masuratorile
            for masurare in data['masuratori_ecografice']:
                row_cells = masuratori_table.add_row().cells
                row_cells[0].text = masurare['structura_anatomica'].capitalize()
                row_cells[1].text = str(masurare['valoare_numerica'])
                row_cells[2].text = masurare['unitate_masura']
        else:
            doc.add_paragraph('Nu au fost identificate masuratori ecografice.')

        # Medicamente
        doc.add_heading('Medicație Prescrisa', level=1)

        if data.get('medicamente'):
            for med in data['medicamente']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{med['nume']}").bold = True
                p.add_run(f" - {med['dozaj']}, {med['frecventa']}")
        else:
            doc.add_paragraph('Nu au fost identificate medicamente.')

        # Simptome
        doc.add_heading('Simptome Raportate', level=1)

        if data.get('simptome'):
            for simptom in data['simptome']:
                doc.add_paragraph(simptom, style='List Bullet')
        else:
            doc.add_paragraph('Nu au fost identificate simptome.')

        # Diagnostice
        doc.add_heading('Diagnostic', level=1)

        if data.get('diagnostice'):
            for diagnostic in data['diagnostice']:
                doc.add_paragraph(diagnostic, style='List Bullet')
        else:
            doc.add_paragraph('Nu au fost identificate diagnostice.')

        # Observații
        if data.get('observatii'):
            doc.add_heading('Observații', level=1)
            for obs in data['observatii']:
                doc.add_paragraph(obs)

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run('___________________________').italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para = doc.add_paragraph('Semnatura medicului')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Salveaza documentul
        doc.save(output_path)
        print(f"✅ Raport Word generat cu succes: {output_path}")

    def create_report_with_template(self, json_path: str, template_path: str, output_path: str):
        """
        Creeaza un raport Word folosind un sablon existent (cu docxtpl)
        sablonul trebuie sa conțina variabile Jinja2
        """
        # incarca datele
        data = self.load_json_data(json_path)

        # incarca sablonul
        doc = DocxTemplate(template_path)

        # Pregateste contextul pentru sablon
        context = {
            'data_raport': datetime.now().strftime("%d.%m.%Y"),
            'ora_raport': datetime.now().strftime("%H:%M"),
            'masuratori': data.get('masuratori_ecografice', []),
            'medicamente': data.get('medicamente', []),
            'simptome': data.get('simptome', []),
            'diagnostice': data.get('diagnostice', []),
            'observatii': data.get('observatii', []),
            'are_masuratori': len(data.get('masuratori_ecografice', [])) > 0,
            'are_medicamente': len(data.get('medicamente', [])) > 0,
            'are_simptome': len(data.get('simptome', [])) > 0,
            'are_diagnostice': len(data.get('diagnostice', [])) > 0
        }

        # Renderizeaza sablonul cu datele
        doc.render(context)

        # Salveaza documentul
        doc.save(output_path)
        print(f"✅ Raport Word generat din sablon: {output_path}")

    def create_template(self, output_path: str = "template_fisa_pacient.docx"):
        """
        Creeaza un sablon Word de baza cu variabile Jinja2
        pentru a fi folosit cu docxtpl
        """
        doc = Document()

        # Titlu
        title = doc.add_heading('FIsA PACIENTULUI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informații generale
        doc.add_heading('Informații Generale', level=1)
        doc.add_paragraph('Data raportului: {{ data_raport }} {{ ora_raport }}')
        doc.add_paragraph('Tip investigație: Ecografie cardiaca')

        # Masuratori
        doc.add_heading('Masuratori Ecografice', level=1)
        doc.add_paragraph('{% if are_masuratori %}')
        doc.add_paragraph('{% for masurare in masuratori %}')
        doc.add_paragraph('• {{ masurare.structura_anatomica }}: {{ masurare.valoare_numerica }} {{ masurare.unitate_masura }}')
        doc.add_paragraph('{% endfor %}')
        doc.add_paragraph('{% else %}')
        doc.add_paragraph('Nu au fost identificate masuratori.')
        doc.add_paragraph('{% endif %}')

        # Medicamente
        doc.add_heading('Medicație', level=1)
        doc.add_paragraph('{% if are_medicamente %}')
        doc.add_paragraph('{% for med in medicamente %}')
        doc.add_paragraph('• {{ med.nume }} - {{ med.dozaj }}, {{ med.frecventa }}')
        doc.add_paragraph('{% endfor %}')
        doc.add_paragraph('{% else %}')
        doc.add_paragraph('Nu au fost identificate medicamente.')
        doc.add_paragraph('{% endif %}')

        # Simptome
        doc.add_heading('Simptome', level=1)
        doc.add_paragraph('{% if are_simptome %}')
        doc.add_paragraph('{% for simptom in simptome %}')
        doc.add_paragraph('• {{ simptom }}')
        doc.add_paragraph('{% endfor %}')
        doc.add_paragraph('{% else %}')
        doc.add_paragraph('Nu au fost raportate simptome.')
        doc.add_paragraph('{% endif %}')

        # Diagnostic
        doc.add_heading('Diagnostic', level=1)
        doc.add_paragraph('{% if are_diagnostice %}')
        doc.add_paragraph('{% for diagnostic in diagnostice %}')
        doc.add_paragraph('• {{ diagnostic }}')
        doc.add_paragraph('{% endfor %}')
        doc.add_paragraph('{% else %}')
        doc.add_paragraph('Nu au fost identificate diagnostice.')
        doc.add_paragraph('{% endif %}')

        # Footer
        doc.add_paragraph()
        doc.add_paragraph('___________________________')
        doc.add_paragraph('Semnatura medicului')

        doc.save(output_path)
        print(f"✅ sablon creat: {output_path}")
        print(f"💡 Poți edita acest sablon in Word si apoi il poți folosi cu create_report_with_template()")


def generate_word_report(json_path: str, output_path: str = None, use_template: bool = False, template_path: str = None):
    """
    Funcție helper pentru generarea rapida a raportului Word

    Args:
        json_path: Calea catre fisierul JSON cu datele
        output_path: Calea unde sa salveze raportul (opțional)
        use_template: Daca True, foloseste un sablon Word
        template_path: Calea catre sablonul Word (necesar daca use_template=True)
    """
    if output_path is None:
        output_path = f"raport_medical_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    generator = MedicalReportGenerator()

    if use_template:
        if template_path is None:
            print("❌ Trebuie sa specifici template_path cand use_template=True")
            return
        generator.create_report_with_template(json_path, template_path, output_path)
    else:
        generator.create_simple_report_without_template(json_path, output_path)

    return output_path


# Exemplu de utilizare
if __name__ == "__main__":
    print("=" * 80)
    print("GENERATOR RAPOARTE MEDICALE WORD")
    print("=" * 80)

    # Exemplu 1: Creeaza un raport simplu (fara sablon)
    print("\n1️⃣  Generare raport simplu (fara sablon)...")
    try:
        generate_word_report(
            json_path="fisa_pacient_medical_structured.json",
            output_path="raport_medical_simplu.docx",
            use_template=False
        )
    except FileNotFoundError:
        print("⚠️  Fisierul JSON nu a fost gasit. Ruleaza mai intai medical_entity_extractor.py")

    # Exemplu 2: Creeaza un sablon Word
    print("\n2️⃣  Generare sablon Word...")
    generator = MedicalReportGenerator()
    generator.create_template("template_fisa_pacient.docx")

    # Exemplu 3: Foloseste sablonul pentru generarea raportului
    print("\n3️⃣  Generare raport folosind sablonul...")
    try:
        generate_word_report(
            json_path="fisa_pacient_medical_structured.json",
            output_path="raport_medical_cu_template.docx",
            use_template=True,
            template_path="template_fisa_pacient.docx"
        )
    except FileNotFoundError:
        print("⚠️  Fisierul JSON sau sablonul nu au fost gasite.")

    print("\n" + "=" * 80)
    print("✅ PROCESARE COMPLETa!")
    print("=" * 80)

